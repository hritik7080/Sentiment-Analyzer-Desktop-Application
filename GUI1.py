from tkinter import *
from tkinter import messagebox as ms
import sqlite3

import sys

try:
    import Tkinter as tk
except ImportError:
    import tkinter as tk

try:
    import ttk

    py3 = False
except ImportError:
    import tkinter.ttk as ttk

    py3 = True
import sys
from keras.models import load_model
from tkinter.filedialog import askopenfile

try:
    import Tkinter as tk
except ImportError:
    import tkinter as tk

try:
    import ttk

    py3 = False
except ImportError:
    import tkinter.ttk as ttk

    py3 = True

import os.path
import pandas as pd
import numpy as np
from sklearn.feature_extraction.text import CountVectorizer
from keras.preprocessing.text import Tokenizer
from keras.preprocessing.sequence import pad_sequences
from keras.models import Sequential
from keras.layers import Dense, Embedding, LSTM, SpatialDropout1D
from sklearn.model_selection import train_test_split
from nltk.stem import WordNetLemmatizer
from keras.utils.np_utils import to_categorical

# LOGIN DATABASE
user=''
with sqlite3.connect('quit.db') as db:
    c = db.cursor()

c.execute('CREATE TABLE IF NOT EXISTS user (username TEXT NOT NULL ,password TEX NOT NULL);')
db.commit()
db.close()


# main Class
class main:
    def __init__(self, master):
        self.master = master
        self.username = StringVar()
        self.password = StringVar()
        self.n_username = StringVar()
        self.n_password = StringVar()
        self.otp=StringVar()
        self.o_username=StringVar()
        self.rotp=0
        self.widgets()
        self.c=0

    def login(self):
        with sqlite3.connect('quit.db') as db:
            c = db.cursor()

        find_user = ('SELECT * FROM user WHERE username = ? and password = ?')
        c.execute(find_user, [(self.username.get()), (self.password.get())])
        result = c.fetchall()
        if result:
            self.logf.pack_forget()
            self.head['text'] = self.username.get() + '\n Logged In'
            self.head['pady'] = 150
            global user
            user=self.username.get()
#%%
            yelp = pd.read_csv('./train.csv', encoding='ISO-8859-1')
            yelp['SentimentText'] = yelp['SentimentText'].apply(lambda x: x.lower())
            lm = WordNetLemmatizer()
            yelp['SentimentText'] = yelp['SentimentText'].apply(lambda x: lm.lemmatize(x))
            import re
            yelp['SentimentText'] = yelp['SentimentText'].apply(lambda x: re.sub('[^a-zA-Z0-9\s]', '', x))
            max_features = 2000
            tokenizer = Tokenizer(num_words=max_features, split=' ')
            tokenizer.fit_on_texts(yelp['SentimentText'].values)
            # %%
            x = tokenizer.texts_to_sequences(yelp['SentimentText'].values)
            # %%
            x = pad_sequences(x)
            y = yelp['Sentiment'].values
            x_train, x_test, y_train, y_test = train_test_split(x, y, test_size=0.33, random_state=42)

            def vp_start_gui():
                '''Starting point when module is the main routine.'''
                global val, w, root
                global prog_location
                prog_call = sys.argv[0]
                prog_location = os.path.split(prog_call)[0]
                root = tk.Tk()
                top = Toplevel1(root)
                root.mainloop()

            w = None

            def create_Toplevel1(root, *args, **kwargs):
                '''Starting point when module is imported by another program.'''
                global w, w_win, rt
                global prog_location
                prog_call = sys.argv[0]
                prog_location = os.path.split(prog_call)[0]
                rt = root
                w = tk.Toplevel(root)
                top = Toplevel1(w)
                return (w, top)

            def destroy_Toplevel1():
                global w
                w.destroy()
                w = None

            class Toplevel1:
                def __init__(self, top=None):
                    '''This class configures and populates the toplevel window.
                       top is the toplevel containing window.'''
                    _bgcolor = '#d9d9d9'  # X11 color: 'gray85'
                    _fgcolor = '#000000'  # X11 color: 'black'
                    _compcolor = '#d9d9d9'  # X11 color: 'gray85'
                    _ana1color = '#d9d9d9'  # X11 color: 'gray85'
                    _ana2color = '#ececec'  # Closest X11 color: 'gray92'
                    font16 = "-family {Segoe UI} -size 14 -weight bold -slant " \
                             "roman -underline 1 -overstrike 0"
                    font17 = "-family {Segoe UI} -size 12 -weight bold -slant " \
                             "roman -underline 0 -overstrike 0"
                    font19 = "-family {Segoe UI} -size 9 -weight bold -slant roman" \
                             " -underline 0 -overstrike 0"
                    font21 = "-family {Segoe UI} -size 10 -weight normal -slant " \
                             "italic -underline 0 -overstrike 0"
                    font22 = "-family {Segoe UI} -size 14 -weight bold -slant " \
                             "roman -underline 0 -overstrike 0"
                    font24 = "-family Verdana -size 18 -weight bold -slant roman " \
                             "-underline 0 -overstrike 0"
                    self.style = ttk.Style()
                    if sys.platform == "win32":
                        self.style.theme_use('winnative')
                    self.style.configure('.', background=_bgcolor)
                    self.style.configure('.', foreground=_fgcolor)
                    self.style.configure('.', font="TkDefaultFont")
                    self.style.map('.', background=
                    [('selected', _compcolor), ('active', _ana2color)])

                    top.geometry("783x830+644+136")
                    top.minsize(148, 1)
                    top.maxsize(1924, 1055)
                    top.resizable(1, 1)
                    top.title("Sentiment Analysis")
                    top.configure(borderwidth="1")
                    top.configure(background="#ffff80")

                    self.Canvas1 = tk.Canvas(top)
                    self.Canvas1.place(relx=0.013, rely=0.012, relheight=0.967
                                       , relwidth=0.962)
                    self.Canvas1.configure(background="#ffff80")
                    self.Canvas1.configure(borderwidth="2")
                    self.Canvas1.configure(insertbackground="black")
                    self.Canvas1.configure(relief="ridge")
                    self.Canvas1.configure(selectbackground="#c4c4c4")
                    self.Canvas1.configure(selectforeground="black")

                    self.Message1 = tk.Message(self.Canvas1)
                    self.Message1.place(relx=0.093, rely=0.012, relheight=0.075
                                        , relwidth=0.831)
                    self.Message1.configure(background="#ffff80")
                    self.Message1.configure(font=font24)
                    self.Message1.configure(foreground="#000080")
                    self.Message1.configure(highlightbackground="#d9d9d9")
                    self.Message1.configure(highlightcolor="#0000ff")
                    self.Message1.configure(text='''**** Sentiment Analyzer ****''')
                    self.Message1.configure(width=626)

                    self.style.configure('TSizegrip', background=_bgcolor)
                    self.TSizegrip1 = ttk.Sizegrip(self.Canvas1)
                    self.TSizegrip1.place(anchor='se', relx=1.0, rely=1.0)

                    self.Labelframe1 = tk.LabelFrame(self.Canvas1)
                    self.Labelframe1.place(relx=0.04, rely=0.075, relheight=0.33
                                           , relwidth=0.93)
                    self.Labelframe1.configure(relief='groove')
                    self.Labelframe1.configure(font=font16)
                    self.Labelframe1.configure(foreground="#000080")
                    self.Labelframe1.configure(text='''Analyzer''')
                    self.Labelframe1.configure(background="#ffff80")
                    self.Labelframe1.configure(highlightbackground="#000080")
                    self.Labelframe1.configure(highlightcolor="#000080")

                    self.Label1 = tk.Label(self.Labelframe1)
                    self.Label1.place(relx=0.029, rely=0.189, height=26, width=222
                                      , bordermode='ignore')
                    self.Label1.configure(background="#ffff80")
                    self.Label1.configure(disabledforeground="#a3a3a3")
                    self.Label1.configure(font=font19)
                    self.Label1.configure(foreground="#000080")
                    self.Label1.configure(text='''Check the sentiment of your review:''')

                    self.Entry1 = tk.Entry(self.Labelframe1)
                    self.Entry1.place(relx=0.043, rely=0.302, height=24, relwidth=0.677
                                      , bordermode='ignore')
                    self.Entry1.configure(background="white")
                    self.Entry1.configure(disabledforeground="#a3a3a3")
                    self.Entry1.configure(font="TkFixedFont")
                    self.Entry1.configure(foreground="#000000")
                    self.Entry1.configure(insertbackground="black")
                    self.Entry1.configure(takefocus="0")

                    self.TSeparator1 = ttk.Separator(self.Labelframe1)
                    self.TSeparator1.place(relx=0.506, rely=0.475, relheight=1.245
                                           , bordermode='ignore')
                    self.TSeparator1.configure(orient="vertical")
                    self.TSeparator1.configure(takefocus="0")

                    self.Label2 = tk.Label(self.Labelframe1)
                    self.Label2.place(relx=0.029, rely=0.642, height=36, width=302
                                      , bordermode='ignore')
                    self.Label2.configure(background="#ffff80")
                    self.Label2.configure(disabledforeground="#a3a3a3")
                    self.Label2.configure(font=font21)
                    self.Label2.configure(foreground="#000080")
                    self.Label2.configure(text='''The sentiment of the entered sentence:''')

                    self.Labelframe2 = tk.LabelFrame(self.Labelframe1)
                    self.Labelframe2.place(relx=0.529, rely=0.528, relheight=0.396
                                           , relwidth=0.443, bordermode='ignore')
                    self.Labelframe2.configure(relief='groove')
                    self.Labelframe2.configure(font=font19)
                    self.Labelframe2.configure(foreground="#000080")
                    self.Labelframe2.configure(text='''Output:''')
                    self.Labelframe2.configure(background="#ffff80")

                    self.Message2 = tk.Message(self.Labelframe2)
                    self.Message2.place(relx=0.032, rely=0.286, relheight=0.571
                                        , relwidth=0.923, bordermode='ignore')
                    self.Message2.configure(background="#ffff80")
                    self.Message2.configure(font=font22)
                    self.Message2.configure(foreground="#000080")
                    self.Message2.configure(highlightbackground="#d9d9d9")
                    self.Message2.configure(highlightcolor="black")
                    self.Message2.configure(text='Sentiment')

                    self.Message2.configure(width=286)

                    def get():
                        x = [self.Entry1.get().lower()]
                        load = load_model('./model1.h5')
                        print(x)
                        x[0] = re.sub('[^a-zA-Z0-9\s]', '', x[0])
                        print(x)
                        text = tokenizer.texts_to_sequences(x)
                        text = pad_sequences(text, maxlen=x_train.shape[1], dtype='int32', value=0)
                        print(text)
                        sentiment = load.predict(text, batch_size=1, verbose=2)
                        print(sentiment)
                        if sentiment > 0.8:
                            y = 'Positive'
                        elif sentiment < 0.3:
                            y = 'Neagtive'
                        else:
                            y = "Neutral"
                        self.Message2.configure(text=y)

                    self.Button1 = tk.Button(self.Labelframe1)
                    self.Button1.place(relx=0.771, rely=0.302, height=23, width=116
                                       , bordermode='ignore')
                    self.Button1.configure(activebackground="#ececec")
                    self.Button1.configure(activeforeground="#000000")
                    self.Button1.configure(background="#000080")
                    self.Button1.configure(disabledforeground="#a3a3a3")
                    self.Button1.configure(font=font19)
                    self.Button1.configure(foreground="#ffffff")
                    self.Button1.configure(highlightbackground="#d9d9d9")
                    self.Button1.configure(highlightcolor="black")
                    self.Button1.configure(pady="0")
                    self.Button1.configure(takefocus="0")
                    self.Button1.configure(text='''Submit''', command=get)
                    self.Labelframe3 = tk.LabelFrame(self.Canvas1)
                    self.Labelframe3.place(relx=0.04, rely=0.411, relheight=0.504
                                           , relwidth=0.93)
                    self.Labelframe3.configure(relief='groove')
                    self.Labelframe3.configure(font=font16)
                    self.Labelframe3.configure(foreground="#000080")
                    self.Labelframe3.configure(text='''Dataset Analyzer''')
                    self.Labelframe3.configure(background="#ffff80")

                    self.Frame1 = tk.Frame(self.Labelframe3)
                    self.Frame1.place(relx=0.014, rely=0.099, relheight=0.877, relwidth=0.979
                                      , bordermode='ignore')
                    self.Frame1.configure(relief='groove')
                    self.Frame1.configure(borderwidth="2")
                    self.Frame1.configure(relief="groove")
                    self.Frame1.configure(background="#ffffff")
                    self.Frame1.configure(highlightbackground="#000080")
                    self.Frame1.configure(highlightcolor="#000080")

                    self.Label3 = tk.Label(self.Frame1)
                    self.Label3.place(relx=0.029, rely=0.056, height=26, width=322)
                    self.Label3.configure(activebackground="#ffffff")
                    self.Label3.configure(background="#ffffff")
                    self.Label3.configure(disabledforeground="#a3a3a3")
                    self.Label3.configure(font=font19)
                    self.Label3.configure(foreground="#000080")
                    self.Label3.configure(text='''Check your dataset(Give your own dataset) :''')

                    def constraints():
                        try:
                            import ttk
                            py3 = False
                        except ImportError:
                            import tkinter.ttk as ttk
                            py3 = True

                        def vp_start_gui():
                            '''Starting point when module is the main routine.'''
                            global val, w, root
                            root = tk.Tk()
                            top = Toplevel1(root)
                            root.mainloop()

                        w = None

                        def create_Toplevel1(root, *args, **kwargs):
                            '''Starting point when module is imported by another program.'''
                            global w, w_win, rt
                            rt = root
                            w = tk.Toplevel(root)
                            top = Toplevel1(w)
                            return (w, top)

                        def destroy_Toplevel1():
                            global w
                            w.destroy()
                            w = None

                        class Toplevel1:
                            def __init__(self, top=None):
                                '''This class configures and populates the toplevel window.
                                   top is the toplevel containing window.'''
                                _bgcolor = '#d9d9d9'  # X11 color: 'gray85'
                                _fgcolor = '#000000'  # X11 color: 'black'
                                _compcolor = '#d9d9d9'  # X11 color: 'gray85'
                                _ana1color = '#d9d9d9'  # X11 color: 'gray85'
                                _ana2color = '#ececec'  # Closest X11 color: 'gray92'
                                self.style = ttk.Style()
                                if sys.platform == "win32":
                                    self.style.theme_use('winnative')
                                self.style.configure('.', background=_bgcolor)
                                self.style.configure('.', foreground=_fgcolor)
                                self.style.map('.', background=
                                [('selected', _compcolor), ('active', _ana2color)])

                                top.geometry("600x181+966+169")
                                top.minsize(148, 1)
                                top.maxsize(1924, 1055)
                                top.resizable(1, 1)
                                top.title("Constraints")
                                top.configure(background="#ffff80")
                                top.configure(highlightbackground="#d9d9d9")
                                top.configure(highlightcolor="black")

                                self.Canvas1 = tk.Canvas(top)
                                self.Canvas1.place(relx=0.017, rely=0.055, relheight=0.901
                                                   , relwidth=0.972)
                                self.Canvas1.configure(background="#ffff80")
                                self.Canvas1.configure(borderwidth="2")
                                self.Canvas1.configure(highlightbackground="#d9d9d9")
                                self.Canvas1.configure(highlightcolor="black")
                                self.Canvas1.configure(insertbackground="black")
                                self.Canvas1.configure(relief="ridge")
                                self.Canvas1.configure(selectbackground="#c4c4c4")
                                self.Canvas1.configure(selectforeground="black")

                                self.style.configure('TSizegrip', background=_bgcolor)
                                self.TSizegrip1 = ttk.Sizegrip(self.Canvas1)
                                self.TSizegrip1.place(anchor='se', relx=1.0, rely=1.0)

                                self.Labelframe1 = tk.LabelFrame(self.Canvas1)
                                self.Labelframe1.place(relx=0.017, rely=0.061, relheight=0.828
                                                       , relwidth=0.961)
                                self.Labelframe1.configure(relief='groove')
                                self.Labelframe1.configure(font="-family {Segoe UI} -size 11 -weight bold")
                                self.Labelframe1.configure(foreground="#000080")
                                self.Labelframe1.configure(text='''Constraints''')
                                self.Labelframe1.configure(background="#ffff80")
                                self.Labelframe1.configure(highlightbackground="#d9d9d9")
                                self.Labelframe1.configure(highlightcolor="black")

                                self.Label2 = tk.Label(self.Labelframe1)
                                self.Label2.place(relx=0.036, rely=0.593, height=26, width=511
                                                  , bordermode='ignore')
                                self.Label2.configure(activebackground="#f9f9f9")
                                self.Label2.configure(activeforeground="black")
                                self.Label2.configure(background="#ffff80")
                                self.Label2.configure(disabledforeground="#a3a3a3")
                                self.Label2.configure(font="-family {Segoe UI} -size 9")
                                self.Label2.configure(foreground="#ff0000")
                                self.Label2.configure(highlightbackground="#d9d9d9")
                                self.Label2.configure(highlightcolor="black")
                                self.Label2.configure(
                                    text='''2. The name of the feature which is having the text should be SentimentText.''')

                                self.Label1 = tk.Label(self.Labelframe1)
                                self.Label1.place(relx=0.018, rely=0.37, height=26, width=362
                                                  , bordermode='ignore')
                                self.Label1.configure(activebackground="#f0f0f0f0f0f0")
                                self.Label1.configure(activeforeground="black")
                                self.Label1.configure(background="#ffff80")
                                self.Label1.configure(disabledforeground="#a3a3a3")
                                self.Label1.configure(font="-family {Segoe UI} -size 9")
                                self.Label1.configure(foreground="#ff0000")
                                self.Label1.configure(highlightbackground="#d9d9d9")
                                self.Label1.configure(highlightcolor="black")
                                self.Label1.configure(text='''1. Format of the file should be comma separated.''')

                        if __name__ == '__main__':
                            vp_start_gui()

                    self.Button2 = tk.Button(self.Frame1)
                    self.Button2.place(relx=0.029, rely=0.141, height=23, width=286)
                    self.Button2.configure(activebackground="#ececec")
                    self.Button2.configure(activeforeground="#000000")
                    self.Button2.configure(background="#000080")
                    self.Button2.configure(disabledforeground="#a3a3a3")
                    self.Button2.configure(font=font19)
                    self.Button2.configure(foreground="#ffffff")
                    self.Button2.configure(highlightbackground="#d9d9d9")
                    self.Button2.configure(highlightcolor="black")
                    self.Button2.configure(pady="0")
                    self.Button2.configure(takefocus="0")
                    self.Button2.configure(text='''Tap to see the constraints of the file''', command=constraints)

                    self.Entry2 = tk.Entry(self.Frame1)
                    self.Entry2.place(relx=0.25, rely=0.225, height=24, relwidth=0.5)
                    self.Entry2.configure(background="white")
                    self.Entry2.configure(borderwidth="4")
                    self.Entry2.configure(disabledforeground="#a3a3a3")
                    self.Entry2.configure(font="TkFixedFont")
                    self.Entry2.configure(foreground="#000000")
                    self.Entry2.configure(highlightbackground="#000080")
                    self.Entry2.configure(highlightcolor="#000080")
                    self.Entry2.configure(insertbackground="black")
                    self.Entry2.configure(takefocus="0")

                    self.Message3 = tk.Message(self.Frame1)
                    self.Message3.place(relx=0.029, rely=0.394, relheight=0.563
                                        , relwidth=0.447)
                    self.Message3.configure(background="#ffffff")
                    self.Message3.configure(foreground="#000080")
                    self.Message3.configure(highlightbackground="#d9d9d9")
                    self.Message3.configure(highlightcolor="black")
                    self.Message3.configure(text=f'Please Wait!! After click analyze button, It may take a while.')

                    self.Message3.configure(width=306)

                    def process(yelp):
                        pd.set_option('display.width', 2000)
                        pd.set_option('display.max_column', 50)
                        print(yelp)

                        yelp['SentimentText'] = yelp['SentimentText'].apply(lambda x: x.lower())
                        import re

                        yelp['SentimentText'] = yelp['SentimentText'].apply(lambda x: re.sub('[^a-zA-Z0-9\s]', '', x))
                        max_features = 2000
                        tokenizer = Tokenizer(num_words=max_features, split=' ')
                        tokenizer.fit_on_texts(yelp['SentimentText'].values)
                        # %%
                        x = tokenizer.texts_to_sequences(yelp['SentimentText'].values)
                        # %%
                        x = pad_sequences(x)
                        print(x)
                        load = load_model('./model1.h5')
                        sentiment = load.predict(x, batch_size=32, verbose=2)
                        print(sentiment.shape)
                        positive = 0
                        negative = 0
                        for x in range(sentiment.shape[0]):
                            if sentiment[x][0] > .5:
                                positive += 1
                            else:
                                negative += 1
                        if positive > negative:
                            message = 'Congratulations, you have more positive reviews.'
                        else:
                            message = 'Sorry, you you have more negative reviews.'
                        self.Message3.configure(
                            text=f'Your dataset is having {yelp.shape[0]} reviews of which {positive} are positive reviews and {negative} are negative. {message}. In the graph there is green color for positive reviews and yellow color for negative reviews. All the details has been sent to you on your email.')
                        import matplotlib.pyplot as plt

                        labels = ['Positive', 'Negative']
                        sizes = [positive, negative]
                        colors = ['yellowgreen', 'pink']
                        patches, texts = plt.pie(sizes, colors=colors, shadow=True, startangle=90)
                        plt.legend(patches, labels, loc="best")
                        plt.axis('equal')
                        plt.tight_layout()
                        plt.savefig('./grpah.png')
                        plt.show()
                        from docx import Document
                        document = Document()

                        p = document.add_paragraph()
                        r = p.add_run()
                        r.add_text(
                            f'Hello, This is a computer generated sentiment report of your dataset.\nYour dataset is having {yelp.shape[0]} reviews of which {positive} are positive reviews and {negative} are negative.\n {message}.')
                        r.add_picture('./grpah.png')
                        r.add_text(
                            'For any assistance, feel free to contact us......\nName: Hritik Gupta, \n Contact: 7080700630, \nEmail:hritikgupta7080@gmail.com')
                        document.save('./demo.docx')
                        from PIL import Image
                        im1 = Image.open("./grpah.png")
                        width = 260
                        height = 200
                        im1 = im1.resize((width, height), Image.NEAREST)
                        im1.save("./grpah.png")
                        self.graph = ttk.Button(self.Frame1)
                        self.graph.place(relx=0.526, rely=0.366, height=220, width=313)
                        self.graph.configure(takefocus="")
                        self.graph.configure(text='''Tbutton''')
                        # photo_location = os.path.join(prog_location, r"E:\\Int213\\Class work\\grpah.png")
                        from PIL import Image, ImageTk

                        print(im1)
                        # global _img0
                        # im2=Image.open('E:\\Int213\\Class work\\grpah.png')
                        # image=ImageTk.PhotoImage(im2)
                        # _img0 = tk.PhotoImage(file=im1)
                        # print(_img0)
                        # self.graph.configure(image=image)
                        # photo=PhotoImage(file=r'E:\Int213\Class work\grpah.png')
                        # self.graph.configure(image=photo)
                        #%%

                        # from PIL import Image, ImageTk
                        #
                        # # load an image with Pillow's [Image]
                        # loaded_image = Image.open('E:\\Int213\\Class work\\grpah.png')
                        #
                        # # convert loaded_image with Pillow's [ImageTK]
                        # PhotoImage(master=self.Frame1, width=313, height=220)
                        # converted_image = ImageTk.PhotoImage(loaded_image)
                        #
                        # # set canvas dimensions from loaded image size
                        # (canvas_width, canvas_height) = loaded_image.size
                        #
                        # # create a Tkinter canvas / window
                        # canvas = Canvas(self.Frame1, width=canvas_width,
                        #                 height=canvas_height)
                        #
                        # # place our converted image in the window
                        # canvas.create_image(0, 0, image=converted_image,
                        #                     anchor=NW)
                        # # I don't understand this bit either
                        # canvas.place(relx=0.526, rely=0.366, height=220, width=313)
                        import tkinter


                        def prop(n):
                            return 360.0 * n / 1000

                        p=(round((positive/yelp.shape[0])*100,2)*1000)/100
                        print(p)
                        c = tkinter.Canvas(self.Frame1, width=220, height=313)
                        c.place(relx=0.526, rely=0.366, height=220, width=313)
                        c.create_arc((2, 2, 152, 152), fill="#00AC36", outline="#FAF402", start=prop(0),
                                     extent=prop(p))
                        c.create_arc((2, 2, 152, 152), fill="#FAF402", outline="#00AC36", start=prop(p),
                                     extent=prop(1000-p))



                        #%%
                        import smtplib
                        from email.mime.multipart import MIMEMultipart
                        from email.mime.text import MIMEText
                        from email.mime.base import MIMEBase
                        from email import encoders

                        fromaddr = "hritikguptapython@gmail.com"
                        toaddr = user

                        msg = MIMEMultipart()
                        msg['From'] = fromaddr
                        msg['To'] = toaddr
                        msg['Subject'] = "Sentiment Report"
                        body = "Hello, This the sentiment report of your dataset. Please go through the attached document."
                        msg.attach(MIMEText(body, 'plain'))
                        filename = "demo.docx"
                        attachment = open("E:\Int213\Class work\demo.docx", "rb")
                        p = MIMEBase('application', 'octet-stream')
                        p.set_payload((attachment).read())
                        encoders.encode_base64(p)
                        p.add_header('Content-Disposition', "attachment; filename= %s" % filename)
                        msg.attach(p)
                        s = smtplib.SMTP('smtp.gmail.com', 587)
                        s.starttls()
                        s.login(fromaddr, "7080700630hritik")
                        text = msg.as_string()
                        s.sendmail(fromaddr, toaddr, text)
                        print("chala gay")
                        s.quit()

                    def get_file():
                        file = askopenfile(mode='r', filetypes=[('CSV File', '*.csv')])

                        self.Entry2.insert(0, file.name)
                        if file is not None:
                            print(file.name)
                            yelp = pd.read_csv(file)
                            self.Button4 = tk.Button(self.Frame1)
                            self.Button4.place(relx=0.79, rely=0.225, height=23, width=116)
                            self.Button4.configure(activebackground="#ececec")
                            self.Button4.configure(activeforeground="#000000")
                            self.Button4.configure(background="#000080")
                            self.Button4.configure(disabledforeground="#a3a3a3")
                            self.Button4.configure(font=font19)
                            self.Button4.configure(foreground="#ffffff")
                            self.Button4.configure(highlightbackground="#d9d9d9")
                            self.Button4.configure(highlightcolor="black")
                            self.Button4.configure(pady="0")
                            self.Button4.configure(takefocus="0")
                            self.Button4.configure(text='''Analyze''')
                            self.Button4.configure(command=lambda: process(yelp))

                    self.Button3 = tk.Button(self.Frame1)
                    self.Button3.place(relx=0.029, rely=0.225, height=23, width=116)
                    self.Button3.configure(activebackground="#ececec")
                    self.Button3.configure(activeforeground="#000000")
                    self.Button3.configure(background="#000080")
                    self.Button3.configure(disabledforeground="#a3a3a3")
                    self.Button3.configure(font=font19)
                    self.Button3.configure(foreground="#ffffff")
                    self.Button3.configure(highlightbackground="#d9d9d9")
                    self.Button3.configure(highlightcolor="black")
                    self.Button3.configure(pady="0")
                    self.Button3.configure(takefocus="0")
                    self.Button3.configure(text='''Browse File''', command=lambda: get_file())

                    self.TSeparator2 = ttk.Separator(self.Frame1)
                    self.TSeparator2.place(relx=0.505, rely=0.355, relheight=1.155)
                    self.TSeparator2.configure(orient="vertical")
                    self.TSeparator2.configure(takefocus="0")

                    def feedback():
                        try:
                            import ttk
                            py3 = False
                        except ImportError:
                            import tkinter.ttk as ttk
                            py3 = True

                        def vp_start_gui():
                            '''Starting point when module is the main routine.'''
                            global val, w, root
                            root = tk.Tk()
                            top = Toplevel1(root)
                            root.mainloop()

                        w = None

                        def create_Toplevel1(root, *args, **kwargs):
                            '''Starting point when module is imported by another program.'''
                            global w, w_win, rt
                            rt = root
                            w = tk.Toplevel(root)
                            top = Toplevel1(w)
                            return (w, top)

                        def destroy_Toplevel1():
                            global w
                            w.destroy()
                            w = None

                        class Toplevel1:
                            def __init__(self, top=None):
                                '''This class configures and populates the toplevel window.
                                   top is the toplevel containing window.'''
                                _bgcolor = '#d9d9d9'  # X11 color: 'gray85'
                                _fgcolor = '#000000'  # X11 color: 'black'
                                _compcolor = '#d9d9d9'  # X11 color: 'gray85'
                                _ana1color = '#d9d9d9'  # X11 color: 'gray85'
                                _ana2color = '#ececec'  # Closest X11 color: 'gray92'
                                font10 = "-family {Segoe UI} -size 15 -weight bold -slant " \
                                         "roman -underline 1 -overstrike 0"
                                font11 = "-family {Segoe UI} -size 10 -weight bold -slant " \
                                         "roman -underline 0 -overstrike 0"
                                font12 = "-family {Segoe UI} -size 9 -weight bold -slant roman" \
                                         " -underline 0 -overstrike 0"
                                self.style = ttk.Style()
                                if sys.platform == "win32":
                                    self.style.theme_use('winnative')
                                self.style.configure('.', background=_bgcolor)
                                self.style.configure('.', foreground=_fgcolor)
                                self.style.map('.', background=
                                [('selected', _compcolor), ('active', _ana2color)])

                                top.geometry("600x322+963+160")
                                top.minsize(148, 1)
                                top.maxsize(1924, 1055)
                                top.resizable(1, 1)
                                top.title("Review&Feedback")
                                top.configure(background="#ffff80")

                                self.Canvas1 = tk.Canvas(top)
                                self.Canvas1.place(relx=0.017, rely=0.031, relheight=0.91
                                                   , relwidth=0.972)
                                self.Canvas1.configure(background="#ffff80")
                                self.Canvas1.configure(borderwidth="2")
                                self.Canvas1.configure(insertbackground="black")
                                self.Canvas1.configure(relief="ridge")
                                self.Canvas1.configure(selectbackground="#c4c4c4")
                                self.Canvas1.configure(selectforeground="black")

                                self.style.configure('TSizegrip', background=_bgcolor)
                                self.TSizegrip1 = ttk.Sizegrip(self.Canvas1)
                                self.TSizegrip1.place(anchor='se', relx=1.0, rely=1.0)

                                self.Labelframe1 = tk.LabelFrame(self.Canvas1)
                                self.Labelframe1.place(relx=0.034, rely=0.034, relheight=0.87
                                                       , relwidth=0.943)
                                self.Labelframe1.configure(relief='groove')
                                self.Labelframe1.configure(font=font10)
                                self.Labelframe1.configure(foreground="#000080")
                                self.Labelframe1.configure(text='''Review and Feedback''')
                                self.Labelframe1.configure(background="#ffff80")

                                self.Entry1 = tk.Entry(self.Labelframe1)
                                self.Entry1.place(relx=0.018, rely=0.314, height=34, relwidth=0.935
                                                  , bordermode='ignore')
                                self.Entry1.configure(background="white")
                                self.Entry1.configure(disabledforeground="#a3a3a3")
                                self.Entry1.configure(font="TkFixedFont")
                                self.Entry1.configure(foreground="#000000")
                                self.Entry1.configure(insertbackground="black")

                                self.Button1 = tk.Button(self.Labelframe1)
                                self.Button1.place(relx=0.364, rely=0.471, height=33, width=136
                                                   , bordermode='ignore')
                                self.Button1.configure(activebackground="#ececec")
                                self.Button1.configure(activeforeground="#000000")
                                self.Button1.configure(background="#000080")
                                self.Button1.configure(disabledforeground="#a3a3a3")
                                self.Button1.configure(font=font11)
                                self.Button1.configure(foreground="#ffffff")
                                self.Button1.configure(highlightbackground="#d9d9d9")
                                self.Button1.configure(highlightcolor="black")
                                self.Button1.configure(pady="0")
                                self.Button1.configure(text='''Submit''')

                                self.Frame1 = tk.Frame(self.Labelframe1)
                                self.Frame1.place(relx=0.036, rely=0.745, relheight=0.176, relwidth=0.918
                                                  , bordermode='ignore')
                                self.Frame1.configure(relief='groove')
                                self.Frame1.configure(borderwidth="4")
                                self.Frame1.configure(relief="groove")
                                self.Frame1.configure(background="#ffff80")

                                self.Label1 = tk.Label(self.Frame1)
                                self.Label1.place(relx=0.02, rely=0.222, height=26, width=222)
                                self.Label1.configure(background="#ffff80")
                                self.Label1.configure(disabledforeground="#a3a3a3")
                                self.Label1.configure(font=font11)
                                self.Label1.configure(foreground="#ff0000")
                                self.Label1.configure(text='''* Please do not abuse here.''')

                                self.Label2 = tk.Label(self.Labelframe1)
                                self.Label2.place(relx=0.018, rely=0.157, height=36, width=312
                                                  , bordermode='ignore')
                                self.Label2.configure(background="#ffff80")
                                self.Label2.configure(disabledforeground="#a3a3a3")
                                self.Label2.configure(font=font12)
                                self.Label2.configure(foreground="#000080")
                                self.Label2.configure(text='''Please share your valuable advice with us:''')

                        if __name__ == '__main__':
                            vp_start_gui()

                    self.Button4 = tk.Button(self.Canvas1)
                    self.Button4.place(relx=0.611, rely=0.934, height=33, width=266)
                    self.Button4.configure(activebackground="#ececec")
                    self.Button4.configure(activeforeground="#000000")
                    self.Button4.configure(background="#000080")
                    self.Button4.configure(disabledforeground="#a3a3a3")
                    self.Button4.configure(font=font17)
                    self.Button4.configure(foreground="#ffffff")
                    self.Button4.configure(highlightbackground="#d9d9d9")
                    self.Button4.configure(highlightcolor="black")
                    self.Button4.configure(pady="0")
                    self.Button4.configure(takefocus="0")
                    self.Button4.configure(text='''Give Us Feedback...''', command=feedback)

                    self.menubar = tk.Menu(top, font="TkMenuFont", bg=_bgcolor, fg=_fgcolor)
                    top.configure(menu=self.menubar)

                @staticmethod
                def popup1(event, *args, **kwargs):
                    Popupmenu1 = tk.Menu(root, tearoff=0)
                    Popupmenu1.configure(activebackground="#f9f9f9")
                    Popupmenu1.configure(activeborderwidth="1")
                    Popupmenu1.configure(activeforeground="black")
                    Popupmenu1.configure(background="#d9d9d9")
                    Popupmenu1.configure(borderwidth="1")
                    Popupmenu1.configure(disabledforeground="#a3a3a3")
                    Popupmenu1.configure(font="{Segoe UI} 9")
                    Popupmenu1.configure(foreground="black")
                    Popupmenu1.configure(tearoff="0")
                    Popupmenu1.post(event.x_root, event.y_root)

            if __name__ == '__main__':
                vp_start_gui()


        else:
            ms.showerror('Oops!', 'Username Not Found.')

    def new_user(self):
        with sqlite3.connect('quit.db') as db:
            c = db.cursor()

        find_user = ('SELECT * FROM user WHERE username = ?')
        c.execute(find_user, [(self.username.get())])
        if c.fetchall():
            ms.showerror('Error!', 'Username Taken Try a Diffrent One.')
        else:
            ms.showinfo('Success!', 'Account Created!')
            self.log()
        insert = 'INSERT INTO user(username,password) VALUES(?,?)'
        c.execute(insert, [(self.n_username.get()), (self.n_password.get())])
        db.commit()
    def forget(self):
        with sqlite3.connect('quit.db') as db:
            c = db.cursor()
        find_user = ('SELECT * FROM user WHERE username = ?')
        c.execute(find_user, [(self.o_username.get())])
        if c.fetchall():
            ms.showinfo('Success','Email Found. OTP has been send to your email.')
            import smtplib
            from email.mime.multipart import MIMEMultipart
            fromaddr="hritikguptapython@gmail.com"
            toaddr=self.o_username.get()
            print(toaddr)
            s=smtplib.SMTP('smtp.gmail.com',587)
            s.starttls()
            s.login(fromaddr,'7080700630hritik')
            msg=MIMEMultipart()

            msg['Subject']='OTP'


            import random
            self.rotp=random.randint(1000,9999)

            s.sendmail(fromaddr,toaddr,f'Your OTP to see your Password is {self.rotp}')


            self.c+=1

        else:
            ms.showerror('Error','Email does nor exist.')
    def log(self):
        self.username.set('')
        self.password.set('')
        self.crf.pack_forget()
        self.head['text'] = 'LOGIN'
        self.logf.pack()
        self.rsf.pack_forget()

    def check_otp(self):
        if self.c==0:
            ms.showerror('Error!', 'Check your Email first.')

        else:
           if self.rotp==int(self.otp.get()):
               with sqlite3.connect('quit.db') as db:
                   c = db.cursor()

               find_pass = ('SELECT password FROM user WHERE username = ?')
               c.execute(find_pass, [(self.o_username.get())])
               ms.showinfo('Success',f'Your Password was "{c.fetchone()[0]}"')
               self.log()


    def cr(self):
        self.n_username.set('')
        self.n_password.set('')
        self.logf.pack_forget()
        self.head['text'] = 'Create Account'
        self.crf.pack()
    def rs(self):
        self.o_username.set('')
        self.otp.set('')
        self.logf.pack_forget()
        self.head['text'] = 'Show Password'
        self.rsf.pack()


    def widgets(self):
        self.head = Label(self.master, text='LOGIN', font=('', 35), pady=10)
        self.head.pack()
        self.logf = Frame(self.master, padx=10, pady=10)
        Label(self.logf, text='Username: ', font=('', 20), pady=5, padx=5).grid(sticky=W)
        Entry(self.logf, textvariable=self.username, bd=5, font=('', 15)).grid(row=0, column=1)
        Label(self.logf, text='Password: ', font=('', 20), pady=5, padx=5).grid(sticky=W)
        Entry(self.logf, textvariable=self.password, bd=5, font=('', 15), show='*').grid(row=1, column=1)
        Button(self.logf, text=' Login ', bd=3, font=('', 15), padx=5, pady=5, command=self.login).grid()
        Button(self.logf, text=' Create Account ', bd=3, font=('', 15), padx=5, pady=5, command=self.cr).grid(row=2,
                                                                                                              column=1)
        Button(self.logf, text=' Forget Password ', bd=3, font=('', 15), padx=5, pady=5, command=self.rs).grid(row=3,
                                                                                                              column=1)
        self.logf.pack()

        self.crf = Frame(self.master, padx=10, pady=10)
        Label(self.crf, text='Username: ', font=('', 20), pady=5, padx=5).grid(sticky=W)
        Entry(self.crf, textvariable=self.n_username, bd=5, font=('', 15)).grid(row=0, column=1)
        Label(self.crf, text='Password: ', font=('', 20), pady=5, padx=5).grid(sticky=W)
        Entry(self.crf, textvariable=self.n_password, bd=5, font=('', 15), show='*').grid(row=1, column=1)
        Button(self.crf, text='Create Account', bd=3, font=('', 15), padx=5, pady=5, command=self.new_user).grid()
        Button(self.crf, text='Go to Login', bd=3, font=('', 15), padx=5, pady=5, command=self.log).grid(row=2,
                                                                                                         column=1)

        self.rsf = Frame(self.master, padx=10, pady=10)
        Label(self.rsf, text='Email: ', font=('', 20), pady=5, padx=5).grid(sticky=W)
        Entry(self.rsf, textvariable=self.o_username, bd=5, font=('', 15)).grid(row=0, column=1)
        Button(self.rsf, text='Check Email', bd=3, font=('', 15), padx=5, pady=5, command=self.forget).grid(row=0,column=2)
        Label(self.rsf, text='OTP: ', font=('', 20), pady=5, padx=5).grid(sticky=W)
        Entry(self.rsf, textvariable=self.otp, bd=5, font=('', 15)).grid(row=1, column=1)
        Button(self.rsf, text='Check OTP', bd=3, font=('', 15), padx=5, pady=5, command=self.check_otp).grid(row=1,
                                                                                                         column=2)



root = Tk()
main(root)
root.mainloop()